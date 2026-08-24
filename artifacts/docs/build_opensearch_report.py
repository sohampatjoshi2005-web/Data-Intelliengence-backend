from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('/Users/sathya/Downloads/Agentic-auto-ml-main 2/artifacts/docs/opensearch_research_and_implementation_report.docx')

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [('Title', 24, RGBColor(12, 74, 110)), ('Heading 1', 15, RGBColor(13, 35, 65)), ('Heading 2', 12, RGBColor(25, 60, 110))]:
        st = styles[name]
        st.font.name = 'Aptos'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('OpenSearch Web Search, Research, and Implementation Report')
    r.font.name = 'Aptos Display'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(12, 74, 110)
    s = doc.add_paragraph()
    s.add_run('Project: Agentic AutoML / Custom GPT Control Plane\n').bold = True
    s.add_run('Scope: Separate OpenSearch web tab, provider chaining, indexing pipeline, and production-readiness summary')
    s.paragraph_format.space_after = Pt(10)
    box = doc.add_table(rows=1, cols=3)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    widths = [2.1, 2.2, 2.4]
    labels = [('Status', 'Working'), ('Providers', 'SearXNG + Tavily'), ('LLM', 'Ollama single-model synthesis')]
    for i, (k, v) in enumerate(labels):
        cell = box.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, 'EAF4FB')
        cell.paragraphs[0].add_run(f'{k}: ').bold = True
        cell.paragraphs[0].add_run(v)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item)


def add_provider_table(doc):
    doc.add_heading('Provider Research Summary', level=1)
    doc.add_paragraph('We evaluated the free/self-hosted path and the managed API path so the OpenSearch web workspace could support both low-cost experimentation and more reliable fallback behavior.')
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['Provider', 'Role', 'Cost Model', 'Strength', 'Current Use']):
        hdr[i].text = text
        set_cell_shading(hdr[i], 'D9EAF7')
    rows = [
        ('SearXNG', 'Primary live search', 'Free / self-hosted', 'No API key, private, controllable', 'Primary when local service is up'),
        ('Tavily', 'Fallback live search', 'Managed API', 'Reliable structured results for AI workflows', 'Fallback when SearXNG is unavailable'),
        ('OpenSearch', 'Index + retrieval store', 'Self-hosted', 'Keyword search, snippets, future hybrid retrieval', 'Stores fetched pages for later search'),
        ('Ollama', 'Answer synthesis', 'Local model', 'Single-model local summarization and context explanation', 'Used to explain indexed search results'),
        ('MCP (future)', 'Tool-access layer', 'Depends on configured server', 'Clean provider abstraction once installed', 'Not configured yet in this workspace'),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_architecture(doc):
    doc.add_heading('Implemented Architecture', level=1)
    doc.add_paragraph('The final implemented path for the dedicated OpenSearch workspace is:')
    flow = doc.add_paragraph()
    flow.style = 'Intense Quote'
    flow.add_run('User Query → Live Search Provider (SearXNG primary, Tavily fallback) → Page Fetch → OpenSearch Index → Indexed Search → Ollama Context Explanation → Clickable Links in Streamlit')
    doc.add_paragraph('Key backend capabilities that were implemented:')
    add_bullets(doc, [
        'Provider abstraction for manual URL ingestion, SearXNG, Tavily, Brave, and SerpAPI.',
        'Provider chaining via environment variables so a self-hosted engine can fail over to a managed API.',
        'Page fetching and content extraction before indexing into OpenSearch.',
        'Context explanation using the same local model already running in Ollama, avoiding a second synthesis model.',
        'Separate OpenSearch tab in Streamlit so web search is isolated from Custom GPT business flows.',
        'Collection-based indexing for repeated searches and future source governance.',
    ])


def add_implementation(doc):
    doc.add_heading('What We Built', level=1)
    doc.add_paragraph('Implementation work completed during the session included both research-driven architecture choices and concrete product changes.')
    doc.add_heading('Backend changes', level=2)
    add_bullets(doc, [
        'Added OpenSearch web service layer with ingest and search endpoints.',
        'Added status endpoint so the UI can show provider readiness and MCP readiness separately.',
        'Added provider chain support: WEB_SEARCH_PROVIDER plus WEB_SEARCH_PROVIDER_FALLBACKS.',
        'Used OpenSearch as the persisted lexical layer for fetched page content and snippets.',
        'Preserved manual URL ingestion so the feature remains useful even without live providers.',
    ])
    doc.add_heading('Frontend changes', level=2)
    add_bullets(doc, [
        'Created a dedicated OpenSearch tab instead of mixing web search into Custom GPT.',
        'Added provider selection, collection naming, URL ingestion, live search, and context explanation output.',
        'Displayed provider used, indexed count, links returned, and explanation output for verification.',
    ])
    doc.add_heading('Operational fixes', level=2)
    add_bullets(doc, [
        'Started and configured OpenSearch successfully.',
        'Brought up SearXNG and patched it to allow JSON responses so the app could consume it programmatically.',
        'Verified Tavily as the immediate working provider path when SearXNG was not yet available.',
        'Kept MCP clearly separated as a future enhancement because no MCP search server was configured in this workspace.',
    ])


def add_findings(doc):
    doc.add_heading('Research Findings and Conclusions', level=1)
    add_numbered(doc, [
        'OpenSearch is not a web search engine by itself. It needs a search/fetch layer in front of it to gather internet content.',
        'SearXNG is the best free/self-hosted choice for the top of the funnel when cost and control matter.',
        'Tavily is a strong fallback because it returns structured results well-suited for AI synthesis workflows.',
        'Using one local Ollama model is sufficient for explanation and context synthesis in the current product stage.',
        'A separate OpenSearch workspace is the right UX decision because web-context workflows differ from CRM/KG business reasoning flows.',
        'MCP should be added only after a real MCP search/fetch server exists; provider APIs were the correct production path for immediate testing.',
    ])


def add_test_results(doc):
    doc.add_heading('Observed Test Results', level=1)
    doc.add_paragraph('The final successful test demonstrated the following working path:')
    add_bullets(doc, [
        'Provider Used: Tavily',
        'Indexed This Run: 8',
        'Links Returned: 7',
        'Context explanation generated successfully by the local model',
        'Clickable links rendered in Streamlit',
    ])
    doc.add_paragraph('Representative live search topics used during testing included:')
    add_bullets(doc, [
        'Neo4j customer 360 graph database examples',
        'Knowledge graph root cause analysis enterprise support',
        'OpenSearch hybrid retrieval graph RAG',
        'Enterprise customer graph churn signals',
    ])


def add_runbook(doc):
    doc.add_heading('Runbook and Required Environment', level=1)
    doc.add_paragraph('To operate the OpenSearch workspace consistently, the following services and environment variables are required.')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['Component', 'Requirement', 'Notes']):
        hdr[i].text = text
        set_cell_shading(hdr[i], 'D9EAF7')
    rows = [
        ('OpenSearch', 'OPENSEARCH_URL + local service', 'Stores fetched pages and supports indexed search'),
        ('SearXNG', 'SEARXNG_BASE_URL + running container/service', 'Primary free provider'),
        ('Tavily', 'TAVILY_API_KEY', 'Fallback provider when SearXNG is down or empty'),
        ('Ollama', 'OLLAMA_OPENAI_BASE_URL + OLLAMA_CHAT_MODEL', 'Single-model explanation layer'),
        ('Backend', 'FastAPI restarted after env changes', 'Provider chain is loaded at runtime startup'),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph('Recommended provider chain for production-like local testing:')
    p = doc.add_paragraph(style='Intense Quote')
    p.add_run('WEB_SEARCH_PROVIDER=searxng | WEB_SEARCH_PROVIDER_FALLBACKS=tavily')


def add_next_steps(doc):
    doc.add_heading('Remaining Work and Next Steps', level=1)
    add_bullets(doc, [
        'Add MCP search/fetch integration once a real MCP server is configured in the workspace.',
        'Persist SearXNG configuration outside the container so JSON support survives container recreation.',
        'Add per-domain controls, crawl refresh, and source governance rules if this becomes a shared team workspace.',
        'Expand answer grounding so the explanation explicitly maps claims to the top returned links.',
        'Optionally add export or briefing generation on top of search results for analyst workflows.',
    ])
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('We researched both free and managed approaches for live web search, selected SearXNG as the preferred free/self-hosted provider, retained Tavily as the practical fallback, and implemented a separate OpenSearch workspace that fetches, indexes, searches, and explains web content using the existing local Ollama model. The feature is now operational for live testing, with MCP intentionally deferred until a real MCP server is available.')


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_provider_table(doc)
    add_architecture(doc)
    add_implementation(doc)
    add_findings(doc)
    add_test_results(doc)
    add_runbook(doc)
    add_next_steps(doc)
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    main()
