from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT_DIR = Path('/Users/sathya/Downloads/Agentic-auto-ml-main 2/artifacts/docs')
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(10.5)
    styles['Heading 1'].font.name = 'Arial'
    styles['Heading 1'].font.size = Pt(18)
    styles['Heading 1'].font.bold = True
    styles['Heading 2'].font.name = 'Arial'
    styles['Heading 2'].font.size = Pt(13)
    styles['Heading 2'].font.bold = True
    styles['Heading 3'].font.name = 'Arial'
    styles['Heading 3'].font.size = Pt(11)
    styles['Heading 3'].font.bold = True


def title(doc: Document, text: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(22)
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True
    r2.font.name = 'Arial'
    r2.font.size = Pt(10.5)


def para(doc: Document, text: str):
    doc.add_paragraph(text)


def bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def numbered(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def footer_note(doc: Document, note: str):
    doc.add_paragraph()
    p = doc.add_paragraph(note)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)


def build_frontend_doc():
    doc = Document()
    set_doc_defaults(doc)
    title(doc, 'Frontend Handoff - OpenSearch Experience', 'UI and integration handoff for the frontend engineer')

    doc.add_heading('Purpose', level=1)
    para(doc, 'This document is the frontend handoff for the OpenSearch experience in the Streamlit control plane. It explains what the frontend is expected to show, what backend endpoints exist, what state should be maintained, and which UX rules should not be broken during future UI work.')

    doc.add_heading('What This Screen Does', level=1)
    bullets(doc, [
        'Acts like a web research chatbot, not a one-shot search form.',
        'Lets the user index URLs into OpenSearch, run live web search, rerank results, and continue asking follow-up questions in the same thread.',
        'Uses automatic model routing across four Ollama models depending on question type.',
        'Keeps episodic, semantic, and procedural memory internally, but memory mechanics should not be displayed to the user.',
        'Shows raw search-engine links first, then final reranked links after retrieval and reranking.'
    ])

    doc.add_heading('Current UI Sections', level=1)
    add_table(doc, ['Section', 'Purpose', 'Notes'], [
        ['Provider / collection controls', 'Choose provider, collection, auto-fetch behavior, thread selection', 'Supports continuing a previous thread or starting a new one.'],
        ['URL ingestion', 'User can paste URLs and fetch plus index them', 'Used for manual workspace-building into OpenSearch.'],
        ['OpenSearch chat', 'Main conversation surface', 'Should feel like a chatbot with a bottom message box after each answer.'],
        ['Latest answer details', 'Shows most recent answer metadata', 'Provider, indexed count, links returned, model used, reranking status.'],
        ['Search Engine Links', 'Raw live-provider links', 'Must be shown before reranked links.'],
        ['Most Relevant Links', 'Final OpenSearch plus reranked links', 'These are the links the answer should primarily rely on.'],
        ['Raw Search Result', 'Debug-only expander', 'Safe to keep for internal/dev use, but not primary UX.']
    ])

    doc.add_heading('Frontend Behavior Rules', level=1)
    numbered(doc, [
        'Keep the chat experience conversational. After every answer, the next message box should remain at the bottom so the user can keep asking follow-up questions.',
        'Do not expose internal memory counts or internal memory object dumps in the main UI.',
        'Always show raw provider links first, then a clearly labeled Most Relevant Links section for reranked results.',
        'Always show which model answered the latest turn.',
        'Do not ask the user to manually choose models for normal chat. Routing is automatic.',
        'Procedural preference updates should happen naturally through chat, not a separate memory settings form.',
        'Thread continuity matters. The same thread should preserve context for follow-up questions.'] )

    doc.add_heading('Model Routing UX', level=1)
    add_table(doc, ['Model', 'Role', 'Typical questions'], [
        ['qwen2.5:3b', 'Fast chat', 'Short follow-ups, quick clarifications, lightweight questions'],
        ['llama3.2:3b', 'Default assistant', 'Normal web research answers and standard synthesis'],
        ['deepseek-r1:1.5b', 'Reasoning / analysis', 'Compare, explain why, analyze tradeoffs, research-style prompts'],
        ['mistral:7b', 'Summarization', 'Summarize, compress, bullets, TL;DR, concise source summaries']
    ])
    para(doc, 'Routing is not purely hardcoded anymore. The system first checks rules, and when those are inconclusive it uses a small Llama-based intent agent to choose the best model among the four.')

    doc.add_heading('Backend Endpoints The Frontend Uses', level=1)
    add_table(doc, ['Endpoint', 'Method', 'Why frontend calls it'], [
        ['/opensearch/status', 'GET', 'Health and provider availability for the OpenSearch workspace'],
        ['/opensearch/web/ingest', 'POST', 'Fetch and index user-provided URLs into OpenSearch'],
        ['/opensearch/web/search', 'POST', 'One-shot search path if needed'],
        ['/opensearch/chat', 'POST', 'Non-streaming chat response path'],
        ['/opensearch/chat/stream', 'POST', 'Streaming chatbot response path used by the current chat UI'],
        ['/opensearch/threads', 'GET', 'List existing conversation threads for the selected collection'],
        ['/opensearch/history', 'GET', 'Load prior thread turns into the visible chat history'],
        ['/opensearch/memory/procedural', 'POST', 'Still exists as an API, but main UX should prefer conversational preference capture']
    ])

    doc.add_heading('What The Streaming Flow Looks Like', level=1)
    numbered(doc, [
        'Frontend sends the chat payload to /opensearch/chat/stream.',
        'Backend emits a model_routing event first when available.',
        'Backend emits answer_chunk events as the chosen LLM streams text.',
        'Backend emits a result event containing the final structured payload.',
        'Frontend updates the assistant bubble live, then stores the final payload in session state for the detail panels below the chat.'
    ])

    doc.add_heading('Important State The Frontend Maintains', level=1)
    bullets(doc, [
        'Active collection name',
        'Current thread id',
        'Visible chat messages',
        'Latest structured result payload for detail panels',
        'Latest ingest result payload',
        'Selected provider and auto-fetch settings',
        'Top K link count'
    ])

    doc.add_heading('Acceptance Checklist For Frontend Changes', level=1)
    bullets(doc, [
        'Chat still supports multiple consecutive follow-up questions in one thread.',
        'Internal memory details are not shown to the user.',
        'Raw provider links appear before reranked links.',
        'Chosen model is visible in the latest answer metadata.',
        'Streaming still works without waiting for the whole answer before rendering.',
        'Switching threads restores prior chat history cleanly.',
        'Starting a new thread clears only thread-local state, not the whole workspace.'
    ])

    footer_note(doc, 'Note: This document reflects the current Streamlit + backend implementation in the repository as of June 5, 2026.')
    path = OUT_DIR / 'frontend_handoff_opensearch.docx'
    doc.save(path)
    return path


def build_custom_gpt_doc():
    doc = Document()
    set_doc_defaults(doc)
    title(doc, 'Custom GPT - Detailed Documentation', 'Architecture, UI flow, routing, connectors, KG, governance, and evaluation')

    doc.add_heading('Purpose', level=1)
    para(doc, 'Custom GPT is the enterprise orchestration surface that combines auth context, connector access, intent routing, structured retrieval, unstructured retrieval, knowledge graph reasoning, hybrid evidence combination, governance, audit, and offline plus production evaluation.')

    doc.add_heading('Primary User Flow', level=1)
    numbered(doc, [
        'Set auth headers and tenant / role context.',
        'Refresh accessible connectors and choose a connector such as CRM or OMS.',
        'Optionally configure connector-specific scope from the streamlined UI controls.',
        'Optionally build an unstructured knowledge base from an uploaded file.',
        'Optionally build a knowledge graph from CRM or OMS into a graph dataset.',
        'Type a natural-language question into the main Query box.',
        'Let the intent router choose structured, unstructured, knowledge_graph, hybrid, or direct_llm.',
        'Inspect answer, route, confidence, execution summary, citations, graph render, and governance / monitoring details.',
        'Run production metrics refresh or DeepEval on recent traces when needed.'
    ])

    doc.add_heading('Top-Level UI Areas', level=1)
    add_table(doc, ['UI area', 'What it covers'], [
        ['Auth Status / headers', 'Tenant, role, JWT context, header preview'],
        ['Connector Access Gateway', 'Refresh catalog, choose connector, configure CRM or OMS filters, validate access'],
        ['Query Surface', 'Main question box plus source toggles and action mode'],
        ['Unstructured Dataset', 'Upload file, choose parser provider, build KB'],
        ['Knowledge Graph', 'Graph dataset id, graph build, graph hops and limits'],
        ['Cypher Workbench', 'Read-only Cypher execution with Graph / Table / Raw result tabs'],
        ['Intent Route (Auto)', 'System-generated route reasoning shown before execution'],
        ['Structured Execution Options', 'Current preview, latest run, connector execution, SQL forcing, debug control'],
        ['Answer + Debug + Eval', 'Answer, rows, SQL, graph, citations, policy, audit, monitoring, metrics, DeepEval']
    ])

    doc.add_heading('Connector Layer', level=1)
    para(doc, 'The connector gateway applies role and tenant aware access checks before Custom GPT can query a source.')
    add_table(doc, ['Connector', 'Current UX shape', 'Important filters'], [
        ['CRM', 'Marketing customer 360 schema', 'view, company, segment, customer_id, email, limit'],
        ['OMS', 'Operational order management schema', 'view, order_number, status, customer_email, limit']
    ])

    doc.add_heading('Supported Routes', level=1)
    add_table(doc, ['Route', 'When it is used', 'What it does'], [
        ['structured', 'Exact facts, analytics, connector-backed table queries, order lookup', 'Runs structured execution and may surface rows and SQL'],
        ['unstructured', 'Document questions over uploaded knowledge bases', 'Queries KB chunks and returns grounded answers with citations'],
        ['knowledge_graph', 'Entity relationships, graph reasoning, connected context', 'Queries graph neighborhood and renders graph output'],
        ['hybrid', 'Multi-source questions that benefit from combined evidence', 'Combines structured, graph, and optionally document evidence'],
        ['direct_llm', 'Fallback for generic non-retrieval questions', 'Uses the LLM directly when retrieval is not needed']
    ])

    doc.add_heading('Intent Routing', level=1)
    para(doc, 'Intent routing happens automatically through /custom-gpt/intent and is also reinforced again inside /custom-gpt/query. The query route applies additional guardrails, such as preferring graph or hybrid for CRM or OMS relationship questions and using hybrid when structured and KG are both active for order or customer-context questions.')

    doc.add_heading('Knowledge Graph Flow', level=1)
    numbered(doc, [
        'Choose Graph Dataset ID.',
        'Build KG from CRM or OMS connector.',
        'Ask a graph-oriented or hybrid question in the main query box.',
        'Custom GPT extracts the best graph seed automatically.',
        'Backend queries the graph and returns graph-aware answer context.',
        'UI auto-renders the graph neighborhood when route is knowledge_graph or hybrid.',
        'Cypher Workbench can be used for read-only graph inspection.'
    ])

    doc.add_heading('Custom GPT Endpoints', level=1)
    add_table(doc, ['Endpoint', 'Method', 'Purpose'], [
        ['/custom-gpt/intent', 'POST', 'Classify route and execution plan before running query'],
        ['/custom-gpt/query', 'POST', 'Execute the final query across structured, unstructured, KG, or hybrid path'],
        ['/custom-gpt/traces', 'GET', 'List audit traces'],
        ['/custom-gpt/metrics', 'GET', 'Return production monitoring summary'],
        ['/custom-gpt/evals/run', 'POST', 'Run eval on recent traces'],
        ['/custom-gpt/evals', 'GET', 'List recent eval runs'],
        ['/kg/build-from-connector', 'POST', 'Build graph from CRM or OMS'],
        ['/kg/subgraph', 'POST', 'Get graph neighborhood for rendering'],
        ['/kg/cypher', 'POST', 'Run safe read-only Cypher'],
        ['/kb/build', 'POST', 'Build unstructured knowledge base'],
        ['/kb/query', 'POST', 'Query unstructured dataset']
    ])

    doc.add_heading('Structured Execution Options Explained', level=1)
    bullets(doc, [
        'Use Current Structured Preview Data: answer from the dataframe currently loaded in the Structured tab.',
        'Use Latest Structured Pipeline Result: answer from the most recent structured pipeline output, model result, or metric summary.',
        'Use Connector Source for Execution: query the selected connector such as CRM or OMS directly.',
        'Force SQL Route: push the structured path toward a stronger SQL-style lookup path.',
        'Include Debug Payload: keep lower-level execution details available in the debug expander.'
    ])

    doc.add_heading('Governance, Audit, and Monitoring', level=1)
    bullets(doc, [
        'Action detection: identifies whether the user is asking the system to do something operational.',
        'Risk classification: labels the answer path with a risk level.',
        'Approval requirement: indicates whether a future human approval step would be required.',
        'Audit trace id: every run can be tied back to a trace.',
        'Monitoring: latency, confidence, citation count, and route-level production metrics are tracked.'
    ])

    doc.add_heading('DeepEval and Offline Evaluation', level=1)
    para(doc, 'The Custom GPT UI includes a dedicated Offline Eval + Production Metrics expander. It can refresh production metrics and run DeepEval on recent traces. The current repo is configured to use Ollama-backed evaluation rather than requiring OpenAI. Faithfulness and answer relevancy are the current default metrics in the UI path.')

    doc.add_heading('What To Give A Frontend or Product Teammate', level=1)
    bullets(doc, [
        'The intended route behavior and which source toggles should be visible.',
        'Connector-specific UI simplification rules for CRM and OMS.',
        'The expectation that KG is driven from the main query box, not separate anchor forms.',
        'The meaning of the answer metadata: route, mode, confidence, SQL, rows, graph, citations, governance, and monitoring.',
        'The fact that debug payloads should remain secondary and not dominate the main user experience.'
    ])

    footer_note(doc, 'This document is intended as a working technical and product reference for the Custom GPT vertical slice in this repository.')
    path = OUT_DIR / 'custom_gpt_detailed_documentation.docx'
    doc.save(path)
    return path


def build_opensearch_doc():
    doc = Document()
    set_doc_defaults(doc)
    title(doc, 'OpenSearch - Detailed Documentation', 'Architecture, chat flow, providers, reranking, memory, and model routing')

    doc.add_heading('Purpose', level=1)
    para(doc, 'The OpenSearch workspace is a web-context chatbot that can search the live web, fetch pages, index them into OpenSearch, rerank hits, answer questions in chat form, and remember thread context across follow-up questions.')

    doc.add_heading('System Flow', level=1)
    numbered(doc, [
        'User asks a question in the OpenSearch chat box.',
        'Provider chain runs live search if enabled, for example SearXNG with Tavily fallback.',
        'Returned links can be fetched and indexed into OpenSearch.',
        'OpenSearch lexical retrieval searches the indexed web collection.',
        'FlashRank reranks the candidate pool.',
        'A model-routing layer chooses the best Ollama model for the question.',
        'Chosen model answers using current hits plus internal memory.',
        'Response streams back to Streamlit while the assistant is still generating.'
    ])

    doc.add_heading('Supported Search Providers', level=1)
    bullets(doc, [
        'manual_urls',
        'searxng',
        'brave',
        'tavily',
        'serpapi'
    ])
    para(doc, 'The provider chain is configurable. Current UX commonly uses SearXNG first with Tavily as fallback.')

    doc.add_heading('OpenSearch Chat UX', level=1)
    add_table(doc, ['Area', 'Purpose'], [
        ['Conversation thread picker', 'Load an existing thread or start a new one'],
        ['URL ingestion area', 'Fetch and index user-provided URLs into the selected collection'],
        ['Chat thread area', 'Shows user and assistant turns'],
        ['Streaming answer area', 'Assistant answer appears progressively during generation'],
        ['Search Engine Links', 'Raw live-provider links before reranking'],
        ['Most Relevant Links', 'Final OpenSearch plus reranked links'],
        ['Model Routing', 'Shows which model answered the latest turn and why'],
        ['Reranking summary', 'Shows reranker backend, model, whether applied, and candidate pool size']
    ])

    doc.add_heading('Model Routing', level=1)
    add_table(doc, ['Model', 'Role', 'Typical trigger'], [
        ['qwen2.5:3b', 'Fast chat', 'Short follow-ups and quick direct questions'],
        ['llama3.2:3b', 'Default assistant', 'Normal web research responses'],
        ['deepseek-r1:1.5b', 'Reasoning / research', 'Analysis, comparison, why, tradeoffs, deeper synthesis'],
        ['mistral:7b', 'Summarization', 'Bullets, TL;DR, concise summaries, compression']
    ])
    para(doc, 'Routing uses both rules and an intent-agent fallback. If rules do not clearly decide, Llama acts as the intent router and chooses among the four models.')

    doc.add_heading('Memory Design', level=1)
    bullets(doc, [
        'Episodic memory: prior turns in the same thread.',
        'Semantic memory: reusable notes distilled from earlier answers.',
        'Procedural memory: answer-style preferences such as always cite links or keep answers concise.',
        'These memories are persisted in OpenSearch and used internally, but should not be surfaced as primary user-facing mechanics.'
    ])

    doc.add_heading('Reranking', level=1)
    para(doc, 'The current active reranking path uses FlashRank instead of the earlier BGE path, because BGE caused native runtime instability on the local Mac environment. The UX requirement is to show raw search-engine links first, then show the reranked final links as the most relevant links.')
    add_table(doc, ['Setting', 'Current value / direction'], [
        ['Backend', 'flashrank'],
        ['Default model', 'ms-marco-MiniLM-L-12-v2'],
        ['Candidate pool', 'Configurable through OPENSEARCH_RERANK_TOP_K'],
        ['Display rule', 'Raw provider links first, reranked links second']
    ])

    doc.add_heading('Important Endpoints', level=1)
    add_table(doc, ['Endpoint', 'Method', 'Purpose'], [
        ['/opensearch/status', 'GET', 'Status and provider readiness'],
        ['/opensearch/web/ingest', 'POST', 'Fetch and index URLs into OpenSearch'],
        ['/opensearch/web/search', 'POST', 'Search indexed web content'],
        ['/opensearch/chat', 'POST', 'Non-streaming chatbot response'],
        ['/opensearch/chat/stream', 'POST', 'Streaming chatbot response'],
        ['/opensearch/threads', 'GET', 'List threads for a collection'],
        ['/opensearch/history', 'GET', 'Fetch episodic thread history'],
        ['/opensearch/memory/procedural', 'POST', 'Store procedural memory directly when needed']
    ])

    doc.add_heading('Streaming Event Contract', level=1)
    bullets(doc, [
        'model_routing event: selected model and routing reason',
        'answer_chunk event: incremental assistant text',
        'result event: final structured payload used by the detail panels',
        'done event: stream completion marker',
        'error event: surfaced when backend fails during the stream'
    ])

    doc.add_heading('Operational Notes', level=1)
    bullets(doc, [
        'OpenSearch must be reachable for lexical indexing and memory persistence.',
        'Ollama must be running for answer generation and intent-agent model routing.',
        'The reranker is optional in principle, but the current UX expects it when enabled.',
        'Conversation continuity depends on thread_id persistence in frontend state.',
        'Procedural preference updates can be expressed naturally in chat, such as From now on, always cite links and keep answers concise.'
    ])

    doc.add_heading('Recommended Questions For Testing', level=1)
    bullets(doc, [
        'What are the main use cases for AI agents in customer support?',
        'Which of those links talks about ticket routing?',
        'Compare rule-based automation with AI agents for support operations.',
        'Summarize the top 3 sources in 4 short bullets.'
    ])

    footer_note(doc, 'This document is meant to stand on its own for engineering, product, and QA conversations about the OpenSearch workspace.')
    path = OUT_DIR / 'opensearch_detailed_documentation.docx'
    doc.save(path)
    return path


if __name__ == '__main__':
    paths = [build_frontend_doc(), build_custom_gpt_doc(), build_opensearch_doc()]
    for p in paths:
        print(p)
