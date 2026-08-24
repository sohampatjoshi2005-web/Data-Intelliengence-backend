from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUT = Path('/Users/sathya/Downloads/Agentic-auto-ml-main 2/artifacts/docs')
OUT.mkdir(parents=True, exist_ok=True)


def setup(doc):
    s = doc.sections[0]
    s.top_margin = Inches(0.65)
    s.bottom_margin = Inches(0.65)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)
    styles = doc.styles
    for name, size, bold in [('Normal', 10.5, False), ('Heading 1', 18, True), ('Heading 2', 13, True), ('Heading 3', 11, True)]:
        st = styles[name]
        st.font.name = 'Arial'
        st.font.size = Pt(size)
        st.font.bold = bold


def title(doc, main, sub):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(main)
    r.font.name = 'Arial'
    r.font.size = Pt(22)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(sub)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    r.italic = True


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def nums(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.italic = True


def custom_gpt_business():
    doc = Document(); setup(doc)
    title(doc, 'Custom GPT Business Documentation', 'Business-facing overview and operating guidance')
    doc.add_heading('What It Is', 1)
    bullets(doc, [
        'Custom GPT is the business-facing orchestration layer for asking natural-language questions across structured data, documents, knowledge graph data, and connector systems.',
        'It helps business users move from source-specific exploration to one guided query experience.',
        'It is designed to support customer, operations, support, and executive use cases with grounded answers and traceability.'
    ])
    doc.add_heading('Business Value', 1)
    bullets(doc, [
        'Reduces switching between CRM, OMS, documents, and graph tools.',
        'Makes connected business reasoning possible, especially for Customer360 and cross-system questions.',
        'Adds governance, confidence, citations, and audit visibility for safer operational use.',
        'Supports future action workflows such as orchestration, campaign operations, and guided follow-up decisions.'
    ])
    doc.add_heading('Typical Business Scenarios', 1)
    table(doc, ['Scenario', 'What Custom GPT does'], [
        ['Customer360', 'Combines customer profile, engagement, intelligence, and journey context into one answer.'],
        ['OMS operations', 'Answers order and customer activity questions using structured and graph routes.'],
        ['Marketing intelligence', 'Surfaces churn risk, CLV, segmentation, next best action, and engagement patterns.'],
        ['Cross-system root cause', 'Uses graph or hybrid reasoning to explain why a customer or issue needs attention.'],
        ['Document-grounded support', 'Retrieves policy or document evidence when a business question depends on written knowledge.'],
    ])
    doc.add_heading('How Business Users Experience It', 1)
    nums(doc, [
        'Choose or validate the connector context.',
        'Optionally build knowledge graph or document datasets when needed.',
        'Ask a question in plain English.',
        'Review route, answer, citations, graph context, and confidence.',
        'Use follow-up questions for clarification or deeper analysis.'
    ])
    doc.add_heading('Governance and Risk', 1)
    bullets(doc, [
        'Custom GPT distinguishes between informational responses and action-like requests.',
        'It records traces and monitoring metrics for each run.',
        'It is designed so higher-risk future actions can require approval rather than running silently.',
        'Business users should treat route, confidence, and citations as part of the answer, not as separate engineering details.'
    ])
    doc.add_heading('Recommended User Guidance', 1)
    bullets(doc, [
        'Ask specific questions with customer, order, or business entities when possible.',
        'Use follow-up prompts to compare, summarize, or explain the answer.',
        'Prefer graph-enabled queries for connected reasoning questions.',
        'Prefer structured routes for exact factual lookups such as order or lead status.'
    ])
    note(doc, 'Business document based on the current Custom GPT implementation in this repository as of June 5, 2026.')
    path = OUT / 'custom_gpt_business_documentation.docx'; doc.save(path); return path


def custom_gpt_technical():
    doc = Document(); setup(doc)
    title(doc, 'Custom GPT Technical Documentation', 'Technical architecture, routing, endpoints, and execution behavior')
    doc.add_heading('Core Architecture', 1)
    bullets(doc, [
        'Custom GPT sits above auth, connector access checks, intent routing, retrieval services, graph services, and observability.',
        'The main execution endpoint is /custom-gpt/query.',
        'Intent classification is handled first, then route-specific services execute structured, unstructured, knowledge_graph, hybrid, or direct_llm behavior.',
        'Final answer shaping adds governance, audit, monitoring, and evaluation support.'
    ])
    doc.add_heading('Key Backend Modules', 1)
    table(doc, ['Module', 'Responsibility'], [
        ['backend/app/services/custom_gpt_intent.py', 'Route classification and source-use decisioning'],
        ['backend/app/services/custom_gpt_structured.py', 'Structured connector-backed execution'],
        ['backend/app/services/custom_gpt_unstructured.py', 'Document / KB execution'],
        ['backend/app/services/custom_gpt_knowledge_graph.py', 'Graph query path'],
        ['backend/app/services/custom_gpt_hybrid.py', 'Combined multi-source execution'],
        ['backend/app/services/custom_gpt_pipeline.py', 'Finalization, governance, and result shaping'],
        ['backend/app/services/custom_gpt_observability.py', 'Trace, metrics, and eval bookkeeping'],
    ])
    doc.add_heading('Routes', 1)
    table(doc, ['Route', 'Trigger pattern', 'Behavior'], [
        ['structured', 'Exact factual or table-oriented questions', 'Runs connector / dataframe / SQL-like structured path'],
        ['unstructured', 'Document knowledge questions', 'Queries knowledge base chunks and produces cited answers'],
        ['knowledge_graph', 'Relationship or connected reasoning questions', 'Queries graph neighborhood and graph answers'],
        ['hybrid', 'Questions needing multiple evidence types', 'Combines structured, docs, and graph evidence'],
        ['direct_llm', 'Fallback general question', 'Uses the LLM directly when retrieval is not needed'],
    ])
    doc.add_heading('Important Endpoints', 1)
    table(doc, ['Endpoint', 'Method', 'Purpose'], [
        ['/custom-gpt/intent', 'POST', 'Return route and source-use guidance'],
        ['/custom-gpt/query', 'POST', 'Run full Custom GPT execution'],
        ['/custom-gpt/traces', 'GET', 'Trace retrieval'],
        ['/custom-gpt/metrics', 'GET', 'Production monitoring summary'],
        ['/custom-gpt/evals/run', 'POST', 'Run eval on recent traces'],
        ['/custom-gpt/evals', 'GET', 'List recent eval runs'],
        ['/kg/build-from-connector', 'POST', 'Build KG from CRM or OMS'],
        ['/kg/cypher', 'POST', 'Read-only graph workbench'],
        ['/kb/build', 'POST', 'Build unstructured knowledge base'],
    ])
    doc.add_heading('Execution Metadata Returned To UI', 1)
    bullets(doc, [
        'Answer', 'Route', 'Mode', 'Confidence', 'Rows or SQL when applicable', 'Execution summary', 'Citations', 'Graph subgraph render data', 'Governance / audit / monitoring info', 'Debug payloads in secondary expanders'
    ])
    doc.add_heading('Evaluation and Monitoring', 1)
    bullets(doc, [
        'Uses production trace summaries for latency, confidence, citation rates, and error monitoring.',
        'Supports DeepEval runs over recent traces.',
        'Current repo uses an Ollama-based path for local evaluation rather than depending on GPT access.'
    ])
    note(doc, 'Technical document aligned to the currently checked-in Custom GPT service and UI modules.')
    path = OUT / 'custom_gpt_technical_documentation.docx'; doc.save(path); return path


def custom_gpt_frontend():
    doc = Document(); setup(doc)
    title(doc, 'Custom GPT Frontend Documentation', 'Frontend behavior, UI sections, and integration contract')
    doc.add_heading('UI Sections', 1)
    table(doc, ['Section', 'Purpose'], [
        ['Auth / header context', 'Displays auth status and request headers used by backend calls'],
        ['Connector Access Gateway', 'Loads connector catalog and configures CRM / OMS scope'],
        ['Query Surface', 'Main question box and source toggles'],
        ['Unstructured Dataset', 'Upload and build document KBs'],
        ['Knowledge Graph', 'Graph dataset id, graph build, graph display settings'],
        ['Cypher Workbench', 'Read-only graph inspection'],
        ['Intent Route', 'Readable route summary before execution'],
        ['Answer section', 'Answer, route, mode, confidence, SQL, rows, graph, citations'],
        ['Offline Eval + Production Metrics', 'Metrics refresh and DeepEval actions'],
    ])
    doc.add_heading('Frontend Rules', 1)
    bullets(doc, [
        'Do not expose raw internal logic before the main answer unless it helps user understanding.',
        'Keep the route and source explanations readable in English, not raw JSON in primary UX.',
        'Graph rendering should be automatic for graph or hybrid answers when a seed entity exists.',
        'Connector-specific UI should stay simplified for CRM and OMS instead of forcing raw JSON input.',
        'Cypher workbench should remain clearly separate from normal user questioning.'
    ])
    doc.add_heading('Frontend To Backend Dependencies', 1)
    bullets(doc, [
        'Depends heavily on /custom-gpt/intent and /custom-gpt/query.',
        'Uses connector access and graph endpoints.',
        'Uses metrics and eval endpoints for the lower monitoring section.',
        'Must preserve auth headers through every API request.'
    ])
    note(doc, 'Frontend handoff for the Custom GPT tab in the Streamlit control plane.')
    path = OUT / 'custom_gpt_frontend_documentation.docx'; doc.save(path); return path


def opensearch_business():
    doc = Document(); setup(doc)
    title(doc, 'OpenSearch Business Documentation', 'Business overview of web-context research and memory workspace')
    doc.add_heading('What It Is', 1)
    bullets(doc, [
        'OpenSearch provides a separate web research workspace inside the product.',
        'It lets users search the web, index useful pages, ask follow-up questions, and keep a conversation thread with context.',
        'It is useful when the answer depends on external knowledge rather than internal CRM, OMS, or uploaded files alone.'
    ])
    doc.add_heading('Business Value', 1)
    bullets(doc, [
        'Creates a research assistant experience instead of a one-shot search page.',
        'Makes external sources easier to compare and summarize.',
        'Supports persistent context across follow-up questions.',
        'Improves answer quality by reranking the most useful links.'
    ])
    doc.add_heading('Good Business Use Cases', 1)
    table(doc, ['Use case', 'Why OpenSearch helps'], [
        ['Competitive research', 'Pulls multiple public sources into one thread and summary.'],
        ['Vendor evaluation', 'Lets teams compare sources and keep follow-up context.'],
        ['Product research', 'Combines raw search links with reranked useful links.'],
        ['External documentation review', 'Useful when company answers depend on public docs or articles.'],
    ])
    doc.add_heading('User Experience Summary', 1)
    nums(doc, [
        'Ask a web research question.',
        'Review the streamed answer.',
        'See raw search-engine links first.',
        'See most relevant reranked links second.',
        'Ask follow-up questions in the same thread.'
    ])
    note(doc, 'Business-facing summary of the OpenSearch workspace as currently implemented.')
    path = OUT / 'opensearch_business_documentation.docx'; doc.save(path); return path


def opensearch_technical():
    doc = Document(); setup(doc)
    title(doc, 'OpenSearch Technical Documentation', 'Provider chain, indexing, reranking, routing, memory, and streaming')
    doc.add_heading('Technical Flow', 1)
    nums(doc, [
        'Live provider search returns candidate URLs.',
        'Pages are fetched and indexed into OpenSearch when enabled.',
        'OpenSearch lexical retrieval finds indexed hits.',
        'FlashRank reranks the candidate pool.',
        'A routing layer selects among four Ollama models.',
        'The chosen model streams the answer back to the UI.',
        'Thread memory is stored in OpenSearch and reused for follow-ups.'
    ])
    doc.add_heading('Model Routing', 1)
    table(doc, ['Model', 'Role'], [
        ['qwen2.5:3b', 'Fast follow-ups and quick chat'],
        ['llama3.2:3b', 'Default assistant and fallback intent chooser'],
        ['deepseek-r1:1.5b', 'Reasoning and research questions'],
        ['mistral:7b', 'Summarization and compression'],
    ])
    doc.add_heading('Important Endpoints', 1)
    table(doc, ['Endpoint', 'Purpose'], [
        ['/opensearch/status', 'Status and provider/reranker/model info'],
        ['/opensearch/web/ingest', 'Fetch and index URLs'],
        ['/opensearch/web/search', 'One-shot web search plus explanation'],
        ['/opensearch/chat', 'Non-streaming chat response'],
        ['/opensearch/chat/stream', 'Streaming chatbot response'],
        ['/opensearch/threads', 'List conversation threads'],
        ['/opensearch/history', 'Load thread history'],
    ])
    doc.add_heading('Memory Model', 1)
    bullets(doc, [
        'Episodic memory stores prior turns in a thread.',
        'Semantic memory stores reusable notes distilled from earlier answers.',
        'Procedural memory stores answer preferences such as always cite links.',
        'These are internal mechanisms and should not dominate the end-user UX.'
    ])
    doc.add_heading('Streaming Contract', 1)
    bullets(doc, [
        'model_routing event', 'answer_chunk event', 'result event', 'done event', 'error event'
    ])
    note(doc, 'Technical documentation for the OpenSearch backend and Streamlit integration currently in this repo.')
    path = OUT / 'opensearch_technical_documentation.docx'; doc.save(path); return path


def opensearch_frontend():
    doc = Document(); setup(doc)
    title(doc, 'OpenSearch Frontend Documentation', 'Frontend behavior and UI contract for the OpenSearch workspace')
    doc.add_heading('Required UX Behavior', 1)
    bullets(doc, [
        'Must behave like a chatbot with persistent follow-up flow.',
        'The message box should remain at the bottom after every assistant answer.',
        'Internal memory should remain hidden from users.',
        'Raw search-engine links should be shown first.',
        'Most relevant reranked links should be shown second under a clear label.',
        'The model used for the latest answer should be visible.',
        'Streaming should render progressively instead of waiting for the full answer.'
    ])
    doc.add_heading('UI Sections', 1)
    table(doc, ['Section', 'Purpose'], [
        ['Workspace controls', 'Collection, provider, auto-fetch settings'],
        ['Thread controls', 'Existing threads and new thread creation'],
        ['URL ingestion', 'Manual URL fetch and indexing'],
        ['Chat thread', 'User and assistant messages'],
        ['Latest answer details', 'Provider, indexed count, links returned, reranking, model routing'],
        ['Search Engine Links', 'Raw provider results'],
        ['Most Relevant Links', 'Final reranked results'],
    ])
    doc.add_heading('Frontend Integration Notes', 1)
    bullets(doc, [
        'Streaming is handled through /opensearch/chat/stream.',
        'Thread selection should rebuild visible chat history.',
        'A new thread should reset chat state cleanly.',
        'Latest answer detail panels should reflect the final result payload, not partial chunks.'
    ])
    note(doc, 'Frontend handoff document for the Streamlit OpenSearch tab.')
    path = OUT / 'opensearch_frontend_documentation.docx'; doc.save(path); return path


if __name__ == '__main__':
    paths = [
        custom_gpt_business(),
        custom_gpt_technical(),
        custom_gpt_frontend(),
        opensearch_business(),
        opensearch_technical(),
        opensearch_frontend(),
    ]
    for p in paths:
        print(p)
